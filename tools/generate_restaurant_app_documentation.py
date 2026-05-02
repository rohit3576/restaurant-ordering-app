from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "Restaurant_QR_App_Documentation.docx"

ACCENT = RGBColor(194, 65, 12)
DARK = RGBColor(15, 23, 42)
MUTED = RGBColor(71, 85, 105)
LIGHT_FILL = "FFF7ED"
HEADER_FILL = "EA580C"
SUBTLE_FILL = "F8FAFC"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(9)
    if color:
        r.font.color.rgb = color
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = True
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for index, text in enumerate(headers):
        set_cell_shading(hdr.cells[index], HEADER_FILL)
        set_cell_text(hdr.cells[index], text, bold=True, color=RGBColor(255, 255, 255))
        if widths:
            hdr.cells[index].width = widths[index]
    for row_values in rows:
        row = table.add_row()
        for index, text in enumerate(row_values):
            set_cell_text(row.cells[index], str(text))
            if index == 0:
                set_cell_shading(row.cells[index], SUBTLE_FILL)
            if widths:
                row.cells[index].width = widths[index]
    doc.add_paragraph()
    return table


def add_bullets(doc, items, style="List Bullet"):
    for item in items:
        doc.add_paragraph(item, style=style)


def add_numbered(doc, items):
    for index, item in enumerate(items, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        r = p.add_run(f"{index}. ")
        r.bold = True
        r.font.size = Pt(9.5)
        r2 = p.add_run(item)
        r2.font.size = Pt(9.5)


def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_FILL)
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = ACCENT
    r.font.size = Pt(10)
    p2 = cell.add_paragraph()
    r2 = p2.add_run(body)
    r2.font.size = Pt(9.5)
    r2.font.color.rgb = MUTED
    doc.add_paragraph()


def add_toc(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "Right-click and update field to refresh the table of contents."
    fld_char3 = OxmlElement("w:fldChar")
    fld_char3.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char)
    run._r.append(instr)
    run._r.append(fld_char2)
    run._r.append(text)
    run._r.append(fld_char3)


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.72)
    section.right_margin = Inches(0.72)

    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10)
    styles["Normal"].paragraph_format.space_after = Pt(6)
    styles["Normal"].paragraph_format.line_spacing = 1.08

    for name, size in [("Title", 26), ("Heading 1", 18), ("Heading 2", 14), ("Heading 3", 11)]:
        style = styles[name]
        style.font.name = "Aptos Display" if name in ["Title", "Heading 1"] else "Aptos"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = ACCENT if name != "Heading 3" else DARK
        style.paragraph_format.space_before = Pt(10 if name != "Title" else 0)
        style.paragraph_format.space_after = Pt(6)

    styles["List Bullet"].font.name = "Aptos"
    styles["List Bullet"].font.size = Pt(9.5)
    styles["List Number"].font.name = "Aptos"
    styles["List Number"].font.size = Pt(9.5)

    settings = doc.settings.element
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    settings.append(update_fields)


def add_title_page(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(28)
    r = p.add_run("QR-Based Restaurant Ordering Web Application")
    r.font.size = Pt(26)
    r.bold = True
    r.font.color.rgb = ACCENT

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Technical Architecture, QA Audit, Feature Documentation, and Free Deployment Guide")
    r.font.size = Pt(13)
    r.font.color.rgb = DARK

    doc.add_paragraph()
    meta = [
        ("Project", "restaurant-ordering-app / Qrave"),
        ("Prepared for", "Client, development team, and potential stakeholders"),
        ("Prepared by", "Senior full-stack engineering and product QA review"),
        ("Repository", str(ROOT)),
        ("Document purpose", "Explain how the system works, what is implemented, what still needs attention, and how to deploy it affordably.")
    ]
    add_table(doc, ["Field", "Details"], meta, [Inches(1.55), Inches(5.6)])

    add_callout(
        doc,
        "Executive note",
        "The application is demo-ready and suitable for stakeholder walkthroughs. For live restaurant operations, the main remaining requirements are persistent production database configuration, secure admin credentials, stronger testing, and operational monitoring."
    )
    doc.add_page_break()


def build_doc():
    doc = Document()
    configure_document(doc)
    add_title_page(doc)

    doc.add_heading("Table of Contents", level=1)
    add_toc(doc)
    doc.add_page_break()

    doc.add_heading("1. Project Overview", level=1)
    doc.add_paragraph(
        "Qrave is a QR-based restaurant ordering web application. Customers scan a table QR code, open a table-specific menu route, browse menu categories, add items to a cart, and place an order. Restaurant staff use an admin/kitchen dashboard to view live orders, update preparation status, manage menu items, and generate printable table QR codes."
    )
    add_table(
        doc,
        ["Area", "Description"],
        [
            ("Project name", "restaurant-ordering-app, branded in the UI as Qrave."),
            ("Purpose", "Reduce waiter dependency for order capture and provide a self-service table ordering workflow."),
            ("Target users", "Restaurant customers, kitchen staff, admins, restaurant owners, and demo stakeholders."),
            ("Core concept", "Each table has a QR code that opens /menu?table=N. The table number is captured and included with every order.")
        ],
        [Inches(1.7), Inches(5.4)]
    )

    doc.add_heading("2. System Architecture", level=1)
    doc.add_paragraph(
        "The project is structured as a React client and an Express API server. It supports two operating modes: a server-backed mode using MongoDB and Socket.io, and a local demo fallback mode using localStorage when the API is unavailable."
    )
    add_table(
        doc,
        ["Layer", "Technology", "Role"],
        [
            ("Frontend", "React 18, Vite, React Router, Tailwind CSS, Framer Motion, Lucide Icons", "Customer menu, cart, checkout, order status, admin dashboard, QR generator, theme toggle."),
            ("API server", "Express, Mongoose, Socket.io, JWT, Helmet, CORS, Compression", "Menu APIs, order APIs, auth, real-time order events, demo store fallback."),
            ("Database", "MongoDB via Mongoose", "Persistent menu, order, and admin data in production."),
            ("Fallback data", "Browser localStorage and in-memory server demo store", "Static/demo operation when backend or MongoDB is not available."),
            ("QR generation", "qrcode npm package", "Creates PNG data URLs for table-specific menu links.")
        ],
        [Inches(1.25), Inches(2.3), Inches(3.55)]
    )
    doc.add_heading("Folder Structure", level=2)
    add_table(
        doc,
        ["Path", "Purpose"],
        [
            ("client/src/pages", "Top-level React screens: CustomerMenu, OrderStatus, AdminLogin, AdminDashboard."),
            ("client/src/components", "Reusable UI components including cart bar, menu cards, checkout modal, QR generator, navigation, and status stepper."),
            ("client/src/services", "API client, Socket.io client, and localStorage fallback store."),
            ("server/src/controllers", "Business logic for auth, menu, orders, and payments."),
            ("server/src/models", "Mongoose schemas for Admin, Menu, and Order."),
            ("server/src/routes", "Express route definitions for API endpoints."),
            ("server/src/data", "Seeded demo menu and demo order store.")
        ],
        [Inches(2.2), Inches(4.9)]
    )

    doc.add_heading("3. Features Breakdown", level=1)
    doc.add_heading("Customer App", level=2)
    add_bullets(doc, [
        "QR routing is implemented through /menu?table=N, with legacy /table/:tableNo still supported.",
        "Table number is normalized and saved to localStorage and sessionStorage.",
        "Menu data loads from /api/menu when available and falls back to local demo menu data.",
        "Categories are derived from menu data and rendered as scrollable tabs.",
        "Cart supports add, increment, decrement, remove-by-zero, total calculation, and per-table localStorage persistence.",
        "Checkout validates non-empty cart and customer name before placing an order.",
        "Order success uses a toast notification with the exact message 'Order placed successfully'.",
        "Order status page shows table, customer, line items, total, and status progression."
    ])
    doc.add_heading("Admin and Kitchen Panel", level=2)
    add_bullets(doc, [
        "Admin login exists with JWT-backed API auth and a local demo fallback for default credentials.",
        "Dashboard displays order count, revenue, in-kitchen count, served count, and an order-flow chart.",
        "Live orders table shows customer, table, items, total, time, and status buttons.",
        "Status flow is pending -> preparing -> served.",
        "Socket.io is used for server-backed real-time order notifications and updates.",
        "Menu manager supports creating, editing, deleting, and availability control for menu items.",
        "A short sound notification is attempted for new server-pushed orders, subject to browser autoplay rules."
    ])
    doc.add_heading("QR Generator", level=2)
    add_bullets(doc, [
        "Admin dashboard includes a 'Generate table QR codes' section.",
        "Input fields support number of tables and base URL, defaulting to window.location.origin.",
        "For each table, the generated URL follows the format {baseUrl}/menu?table=N.",
        "Each table QR displays a table label, 'Scan to Order' text, the encoded URL, and a PNG download button.",
        "Print support is implemented through a print-optimized grid layout.",
        "A 'Scan Me Preview' section shows the first generated table QR code."
    ])

    doc.add_heading("4. Workflow Explanation", level=1)
    add_numbered(doc, [
        "Restaurant staff generate QR codes in the admin dashboard and print one code per table.",
        "A customer scans a table QR. The browser opens /menu?table=N, where N is the assigned table number.",
        "The customer browses categories, adds items to cart, changes quantities, and reviews the total.",
        "The checkout modal collects the customer name and submits the order with table number and line items.",
        "The backend validates menu item availability, calculates total, prevents very recent duplicate submissions, stores the order, and emits a Socket.io event.",
        "The admin dashboard receives the order and staff move it through pending, preparing, and served.",
        "The customer can follow the order status page for updates."
    ])

    doc.add_heading("5. Code Quality Review", level=1)
    add_table(
        doc,
        ["Category", "Assessment"],
        [
            ("Structure", "Clear client/server split with pages, components, services, controllers, routes, and models."),
            ("Modularity", "Good reuse of UI components and service wrappers. QR generation and local store are separated cleanly."),
            ("Naming", "Mostly readable names. One minor consistency issue remains: tableNo and table both exist in fallback order shape."),
            ("Maintainability score", "7.8 / 10 for a demo-ready app. It is organized and understandable, but needs tests, stronger production auth, and better backend operational hardening."),
            ("Developer experience", "Root scripts are present. Vercel and Netlify SPA rewrite files are included.")
        ],
        [Inches(1.7), Inches(5.4)]
    )
    doc.add_heading("Positive Practices", level=2)
    add_bullets(doc, [
        "API calls are centralized through an Axios service.",
        "Socket logic is isolated from page components.",
        "The UI uses reusable components and consistent Tailwind utility patterns.",
        "Server controllers validate important inputs such as empty orders and invalid statuses.",
        "Deployment routing is handled through platform-specific rewrite files."
    ])
    doc.add_heading("Anti-Patterns and Risks", level=2)
    add_bullets(doc, [
        "No automated test suite is present for customer flows, admin flows, API contracts, or QR generation.",
        "Default admin credentials are demo-friendly but unsafe for production if not changed.",
        "The localStorage fallback is excellent for demos but not a multi-device source of truth.",
        "Some image assets depend on external URLs, so menu visuals rely on third-party availability.",
        "Admin destructive actions, such as menu deletion, do not require confirmation.",
        "Payment integration exists as server code but is not presented as a completed customer checkout flow."
    ])

    doc.add_heading("6. Bugs and Issues Found", level=1)
    add_table(
        doc,
        ["Severity", "Issue", "Impact", "Recommendation"],
        [
            ("High", "Production admin credentials must be changed.", "Unauthorized access risk if demo defaults are reused.", "Set strong credentials and protect admin creation/reset workflow."),
            ("High", "No automated tests.", "Regression risk in order placement, QR routing, and admin status updates.", "Add Playwright end-to-end tests and API tests."),
            ("Medium", "localStorage fallback is not shared across devices.", "Static-only deployment cannot support real kitchen devices seeing customer orders.", "Use the backend with MongoDB or Firebase for live operations."),
            ("Medium", "Menu delete has no confirmation dialog.", "Accidental data loss in admin panel.", "Add confirm modal or soft-delete status."),
            ("Medium", "External image URLs can fail or slow the menu.", "Broken visuals and inconsistent performance.", "Upload optimized restaurant-owned images or use a CDN."),
            ("Medium", "Socket URL and API URL must be configured per environment.", "Admin real-time updates fail if env variables point to the wrong host.", "Set VITE_API_URL, VITE_SOCKET_URL, CLIENT_URL, and CORS allowlist correctly."),
            ("Low", "QRCode 'download all' is implemented as print grid, not zip.", "Acceptable for operations, but not a true bulk PNG export.", "Add JSZip if bulk downloadable PNG files are required."),
            ("Low", "Sound notification can be blocked until user interaction.", "New order sound may not play immediately in some browsers.", "Show visual notification and request interaction/permission where appropriate.")
        ],
        [Inches(0.75), Inches(2.0), Inches(2.1), Inches(2.25)]
    )

    doc.add_heading("7. Improvements and Recommendations", level=1)
    doc.add_heading("UI and UX", level=2)
    add_bullets(doc, [
        "Add restaurant branding controls for logo, colors, opening hours, and service notices.",
        "Add item modifiers such as spice level, add-ons, notes, and variants.",
        "Add a clearer empty-cart state and optional order review screen before final placement.",
        "Add table-specific service actions such as 'Call waiter' and 'Request bill'."
    ])
    doc.add_heading("Backend and Operations", level=2)
    add_bullets(doc, [
        "Move from demo data to managed MongoDB Atlas for production.",
        "Add rate limiting, request validation schemas, and structured server logging.",
        "Add admin role permissions for owner, manager, and kitchen staff.",
        "Add audit logs for status updates and menu changes.",
        "Add backup/export capability for orders and menu data."
    ])
    doc.add_heading("SaaS Scalability", level=2)
    add_bullets(doc, [
        "Introduce tenant/restaurant IDs across menus, orders, admins, and QR codes.",
        "Support multiple branches, tables per branch, subscription plans, and branded customer URLs.",
        "Add analytics for average order value, peak hours, best-selling items, kitchen time, and repeat customers.",
        "Separate public customer API from internal admin API and enforce stricter access controls."
    ])

    doc.add_heading("8. Free Deployment Guide", level=1)
    add_callout(
        doc,
        "Beginner-friendly summary",
        "Deploy the frontend to Vercel or Netlify for free. Deploy the backend to Render or Railway free tier. Use MongoDB Atlas free tier for the database. After deployment, regenerate QR codes using the live frontend URL."
    )
    doc.add_heading("Frontend Deployment on Vercel", level=2)
    add_numbered(doc, [
        "Create a GitHub account and push the project repository to GitHub.",
        "Go to vercel.com and sign in with GitHub.",
        "Click Add New Project and import the repository.",
        "Set the Root Directory to client.",
        "Use npm install as the install command and npm run build as the build command.",
        "Use dist as the output directory.",
        "Add environment variables when using a backend: VITE_API_URL=https://your-api.onrender.com/api and VITE_SOCKET_URL=https://your-api.onrender.com.",
        "Click Deploy. Vercel provides a free .vercel.app URL."
    ])
    doc.add_heading("Frontend Deployment on Netlify", level=2)
    add_numbered(doc, [
        "Push the repository to GitHub.",
        "Go to netlify.com and create a new site from Git.",
        "Select the repository and set Base directory to client.",
        "Set Build command to npm run build and Publish directory to client/dist.",
        "Add the same VITE_API_URL and VITE_SOCKET_URL environment variables if the backend is deployed.",
        "Deploy the site. The included _redirects file supports React Router refreshes."
    ])
    doc.add_heading("Backend Deployment on Render", level=2)
    add_numbered(doc, [
        "Create a free MongoDB Atlas cluster and copy the connection string.",
        "Go to render.com and create a new Web Service from the GitHub repository.",
        "Set Root Directory to server.",
        "Set Build Command to npm install.",
        "Set Start Command to npm start.",
        "Add environment variables: MONGO_URI, JWT_SECRET, CLIENT_URL, PORT if needed, and ALLOW_DEMO_STORE=false for production.",
        "Deploy and copy the backend URL, for example https://your-api.onrender.com.",
        "Return to Vercel or Netlify and set VITE_API_URL and VITE_SOCKET_URL using the backend URL."
    ])
    doc.add_heading("Backend Deployment on Railway", level=2)
    add_bullets(doc, [
        "Railway can deploy the server directory from GitHub and provide environment variables through its dashboard.",
        "Use the same required variables: MONGO_URI, JWT_SECRET, CLIENT_URL, and optionally ALLOW_DEMO_STORE=false.",
        "Free availability can change over time, so confirm the current free limits before using it for a client."
    ])
    doc.add_heading("Domain Setup", level=2)
    add_bullets(doc, [
        "Free subdomains are included by Vercel, Netlify, Render, and Railway.",
        "For a custom domain, buy a domain or use a free/low-cost provider, then add the domain inside Vercel or Netlify.",
        "Update DNS records exactly as the hosting platform instructs.",
        "After the custom domain works, regenerate all QR codes with the new public URL."
    ])
    doc.add_heading("QR Code Update After Deployment", level=2)
    add_numbered(doc, [
        "Open the deployed admin dashboard.",
        "Go to Generate table QR codes.",
        "Set Base URL to the live frontend domain, not localhost.",
        "Enter the number of tables.",
        "Download or print the generated table QR codes.",
        "Physically test at least one QR code for table 1, one middle table, and the last table."
    ])
    doc.add_heading("Common Deployment Errors", level=2)
    add_table(
        doc,
        ["Problem", "Likely Cause", "Fix"],
        [
            ("Images not loading", "External URL blocked, removed, or slow.", "Use reliable CDN-hosted images or upload restaurant-owned images."),
            ("API not working", "Wrong VITE_API_URL or CORS CLIENT_URL missing.", "Check frontend env vars and backend allowed origins."),
            ("Admin live updates not working", "Wrong VITE_SOCKET_URL or WebSocket blocked.", "Set socket origin to backend domain and confirm Render/Railway supports WebSockets."),
            ("Refresh gives 404", "SPA rewrite missing.", "Use client/vercel.json on Vercel or client/public/_redirects on Netlify."),
            ("Orders disappear", "Running only static localStorage demo mode.", "Deploy backend and MongoDB for real multi-device persistence."),
            ("Mongo connection fails", "Incorrect MONGO_URI or IP/network restrictions.", "Use MongoDB Atlas connection string and correct database user password.")
        ],
        [Inches(1.8), Inches(2.25), Inches(3.0)]
    )

    doc.add_heading("9. Deployment Readiness Check", level=1)
    add_table(
        doc,
        ["Area", "Current Status", "Before Live Use"],
        [
            ("Customer ordering", "Demo-ready and functional.", "Run real-device QA across common mobile browsers."),
            ("Admin dashboard", "Demo-ready with live order management.", "Secure credentials and add role-based access."),
            ("QR generator", "Implemented with PNG and print workflow.", "Regenerate using deployed URL and test printed codes."),
            ("Persistence", "MongoDB supported; local fallback exists.", "Configure production MongoDB Atlas and disable demo fallback."),
            ("Security", "Basic JWT auth and Helmet are present.", "Change defaults, add rate limits, validate all inputs, and protect admin operations."),
            ("Testing", "Manual QA performed; no automated tests found.", "Add automated regression tests before live restaurant rollout.")
        ],
        [Inches(1.6), Inches(2.5), Inches(3.0)]
    )
    add_callout(
        doc,
        "Readiness verdict",
        "The app is ready for demos, investor walkthroughs, and controlled pilots. It should not be treated as fully production-hardened until database, credentials, monitoring, backups, validation, and automated tests are completed."
    )

    doc.add_heading("10. Future Roadmap", level=1)
    add_table(
        doc,
        ["Phase", "Roadmap Items"],
        [
            ("Phase 1: Live pilot", "Production MongoDB, secure admin, real restaurant menu, printed QR codes, staff training, mobile QA."),
            ("Phase 2: Operational polish", "Order notes, item modifiers, waiter calls, bill request, kitchen filters, order timers, receipt export."),
            ("Phase 3: Payments", "Razorpay or Stripe checkout, payment status, refunds, settlement reporting."),
            ("Phase 4: SaaS", "Multi-restaurant tenancy, branch management, subscriptions, owner analytics, branded domains."),
            ("Phase 5: Advanced analytics", "Sales dashboards, item performance, table turnover, peak hour predictions, inventory signals.")
        ],
        [Inches(1.7), Inches(5.4)]
    )

    doc.add_heading("Appendix A: Verified Implementation Notes", level=1)
    add_bullets(doc, [
        "Browser QA verified /menu?table=3 table detection, cart persistence after refresh, order placement, admin order visibility, QR code URL generation, /menu?table=7 routing, and no observed console errors during those flows.",
        "Production client build completed successfully with npm run build --prefix client.",
        "Server syntax checks passed for the modified order controller and Order model.",
        "The app currently running in the in-app browser is at http://127.0.0.1:5173/admin/dashboard."
    ])

    doc.add_heading("Appendix B: Regenerating This Document", level=1)
    doc.add_paragraph("From the repository root, run:")
    p = doc.add_paragraph()
    r = p.add_run(
        '"C:\\Users\\ashok\\.cache\\codex-runtimes\\codex-primary-runtime\\dependencies\\python\\python.exe" tools\\generate_restaurant_app_documentation.py'
    )
    r.font.name = "Consolas"
    r.font.size = Pt(9)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)


if __name__ == "__main__":
    build_doc()
    print(OUT)
